from __future__ import annotations

import json
import os
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4
from zipfile import ZIP_DEFLATED, ZipFile

EVIA_COST = "150 000 ₽"
MISSING = "MISSING/НЕ ПРЕДОСТАВЛЕНО"
PIPELINE_ORDER = ["0", "1", "2", "3", "3.5", "3.6", "3.7", "4", "5", "6"]
PROMPT_FILE_BY_STAGE = {
    "0": "prompt_0.md",
    "1": "prompt_1.md",
    "2": "prompt_2.md",
    "3": "prompt_3.md",
    "3.5": "prompt_3_5.md",
    "3.6": "prompt_3_6.md",
    "3.7": "prompt_3_7.md",
    "4": "prompt_4.md",
    "5": "prompt_5.md",
    "6": "prompt_6.md",
}


@dataclass
class StageOutput:
    stage: str
    markdown: str
    data: dict[str, Any]


def _sanitize_value(value: Any) -> Any:
    if value is None:
        return MISSING
    if isinstance(value, str):
        cleaned = value.strip()
        return cleaned if cleaned else MISSING
    if isinstance(value, list):
        return [_sanitize_value(v) for v in value] or [MISSING]
    if isinstance(value, dict):
        return {k: _sanitize_value(v) for k, v in value.items()}
    return value


def _prepare_intake(intake: dict[str, Any]) -> dict[str, Any]:
    return {k: _sanitize_value(v) for k, v in intake.items()}


def _load_prompt_text(stage: str) -> str:
    prompt_path = Path("prompts") / PROMPT_FILE_BY_STAGE[stage]
    return prompt_path.read_text(encoding="utf-8")


def _prompt0_gate(intake: dict[str, Any]) -> tuple[str, list[str]]:
    required = ["project_name", "company_name", "goals", "problem_statement"]
    missing = [field for field in required if intake.get(field, MISSING) == MISSING]
    if missing:
        questions = [f"Уточните поле '{field}' для продолжения." for field in missing[:7]]
        return "НЕ ГОТОВО", questions
    return "ГОТОВО", []


def _fixation_block(run_id: str, intake: dict[str, Any], decision: str) -> dict[str, str]:
    return {
        "run_id": run_id,
        "evia_cost_45_days": EVIA_COST,
        "evia_decision": decision,
        "project_name": str(intake.get("project_name", MISSING)),
        "company_name": str(intake.get("company_name", MISSING)),
        "goals": str(intake.get("goals", MISSING)),
        "problem_statement": str(intake.get("problem_statement", MISSING)),
    }


def _dry_run_enabled() -> bool:
    return os.getenv("DRY_RUN", "true").lower() == "true"


def _call_llm(stage: str, prompt_text: str, context: dict[str, Any]) -> str:
    if _dry_run_enabled():
        return (
            f"[DRY_RUN] Stage {stage}\n"
            f"Краткий результат без внешнего LLM.\n"
            f"Контекст содержит поля: {', '.join(sorted(context.keys()))}."
        )

    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is required when DRY_RUN=false")

    model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    user_payload = {
        "stage": stage,
        "prompt": prompt_text,
        "context": context,
        "constraints": {
            "evia_cost_45_days": EVIA_COST,
            "missing_marker": MISSING,
            "no_external_data": True,
        },
    }

    body = json.dumps(
        {
            "model": model,
            "temperature": 0,
            "messages": [
                {
                    "role": "system",
                    "content": "Ты формируешь текст stage-вывода строго по входному prompt и контексту.",
                },
                {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
            ],
        }
    ).encode("utf-8")

    req = urllib.request.Request(
        "https://api.openai.com/v1/chat/completions",
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=60) as response:  # noqa: S310
        payload = json.loads(response.read().decode("utf-8"))

    return payload["choices"][0]["message"]["content"].strip()




def _build_s6_text_from_fixation(copied: dict[str, str], status: str) -> str:
    return (
        "Финальный результат EVIA R&D\n\n"
        f"status: {status}\n"
        f"run_id: {copied['run_id']}\n"
        f"project_name: {copied['project_name']}\n"
        f"company_name: {copied['company_name']}\n"
        f"goals: {copied['goals']}\n"
        f"problem_statement: {copied['problem_statement']}\n"
        f"evia_decision: {copied['evia_decision']}\n"
        f"evia_cost_45_days: {copied['evia_cost_45_days']}"
    )

def _render_markdown(stage: str, prompt_text: str, llm_output: str, data: dict[str, Any]) -> str:
    return (
        f"# S{stage}\n\n"
        "## Prompt source\n\n"
        f"```md\n{prompt_text}\n```\n\n"
        "## LLM output\n\n"
        f"{llm_output}\n\n"
        "## Metadata\n\n"
        f"```json\n{json.dumps(data, ensure_ascii=False, indent=2)}\n```\n"
    )


def _create_docx_from_text(target: Path, title: str, body: str) -> None:
    try:
        from docx import Document

        document = Document()
        document.add_heading(title, level=1)
        for line in body.splitlines():
            document.add_paragraph(line)
        document.save(target)
    except ModuleNotFoundError:
        target.with_suffix(".md").write_text(f"# {title}\n\n{body}", encoding="utf-8")


def generate_run(intake_raw: dict[str, Any], evidence: list[dict[str, str]], root: Path) -> tuple[str, Path]:
    run_id = uuid4().hex[:12]
    run_dir = root / run_id
    run_dir.mkdir(parents=True, exist_ok=True)

    intake = _prepare_intake(intake_raw)
    stage_outputs: list[StageOutput] = []
    stage_context: dict[str, Any] = {"run_id": run_id, "intake": intake, "evidence": evidence}

    s0_prompt = _load_prompt_text("0")
    status, questions = _prompt0_gate(intake)
    s0_llm = _call_llm("0", s0_prompt, {**stage_context, "status": status, "questions": questions})
    s0_data = {
        "stage": "0",
        "status": status,
        "questions": questions,
        "intake": intake,
        "evidence": evidence,
        "auto_web": "disabled",
        "llm_output": s0_llm,
        "dry_run": _dry_run_enabled(),
    }
    stage_outputs.append(StageOutput("0", _render_markdown("0", s0_prompt, s0_llm, s0_data), s0_data))

    if status == "НЕ ГОТОВО":
        _persist_outputs(stage_outputs, run_dir)
        return run_id, _package_outputs(run_id, run_dir)

    decision = "БЕРЁМ" if EVIA_COST == "150 000 ₽" else "НЕ БЕРЁМ"
    if decision == "НЕ БЕРЁМ":
        stage_context["failure_reason"] = "EVIA_COST mismatch"

    for stage in PIPELINE_ORDER[1:8]:
        prompt_text = _load_prompt_text(stage)
        context = {
            **stage_context,
            "stage": stage,
            "evia_cost_45_days": EVIA_COST,
            "evia_decision": decision,
            "previous_stage": stage_outputs[-1].data,
        }
        llm_output = _call_llm(stage, prompt_text, context)
        data = {
            "stage": stage,
            "run_id": run_id,
            "evia_cost_45_days": EVIA_COST,
            "evia_decision": decision,
            "llm_output": llm_output,
        }
        stage_outputs.append(StageOutput(stage, _render_markdown(stage, prompt_text, llm_output, data), data))

    fixation = _fixation_block(run_id, intake, decision)
    s5_prompt = _load_prompt_text("5")
    s5_llm = _call_llm("5", s5_prompt, {**stage_context, "fixation": fixation})
    s5_data = {
        "stage": "5",
        "run_id": run_id,
        "evia_cost_45_days": EVIA_COST,
        "llm_output": s5_llm,
        "(10) БЛОК ФИКСАЦИИ": fixation,
    }
    s5_md = _render_markdown("5", s5_prompt, s5_llm, s5_data) + "\n(10) БЛОК ФИКСАЦИИ\n" + "\n".join(
        f"{k}: {v}" for k, v in fixation.items()
    )
    stage_outputs.append(StageOutput("5", s5_md, s5_data))

    required_fixation_fields = [
        "run_id",
        "evia_cost_45_days",
        "evia_decision",
        "project_name",
        "company_name",
        "goals",
        "problem_statement",
    ]
    missing_fixation = [key for key in required_fixation_fields if key not in fixation]
    if missing_fixation:
        s6_data = {
            "stage": "6",
            "run_id": run_id,
            "status": "OUTPUT INVALID",
            "missing_fixation_fields": missing_fixation,
            "llm_output": "OUTPUT INVALID",
        }
        s6_text = "OUTPUT INVALID"
    else:
        copied = {k: fixation[k] for k in required_fixation_fields}
        s6_text = _build_s6_text_from_fixation(copied, "OK")
        s6_data = {
            "stage": "6",
            "run_id": run_id,
            "status": "OK",
            "copied_from_fixation": copied,
            "evia_cost_45_days": copied["evia_cost_45_days"],
            "evia_decision": copied["evia_decision"],
            "llm_output": s6_text,
        }

    stage_outputs.append(
        StageOutput("6", _render_markdown("6", _load_prompt_text("6"), s6_text, s6_data), s6_data)
    )

    _persist_outputs(stage_outputs, run_dir)
    _create_docx_from_text(run_dir / "S6.docx", "S6 Final", s6_text)
    return run_id, _package_outputs(run_id, run_dir)


def _persist_outputs(outputs: list[StageOutput], run_dir: Path) -> None:
    for output in outputs:
        stage_tag = output.stage.replace(".", "_")
        (run_dir / f"S{stage_tag}.md").write_text(output.markdown, encoding="utf-8")
        (run_dir / f"S{stage_tag}.json").write_text(
            json.dumps(output.data, ensure_ascii=False, indent=2), encoding="utf-8"
        )


def _package_outputs(run_id: str, run_dir: Path) -> Path:
    zip_path = run_dir / f"{run_id}.zip"
    with ZipFile(zip_path, "w", compression=ZIP_DEFLATED) as zf:
        for path in sorted(run_dir.iterdir()):
            if path.name.endswith(".zip"):
                continue
            zf.write(path, arcname=path.name)
    return zip_path
